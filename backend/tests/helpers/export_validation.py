"""Export document validation helpers for integration tests.

Provides utilities to validate exported documents (DOCX, PDF, Markdown)
contain expected citations and content.

Usage:
    from tests.helpers import validate_docx_citations, validate_pdf_citations

    # Validate DOCX citations
    assert validate_docx_citations(docx_bytes, ["Source 1", "Source 2"])

    # Validate PDF citations
    assert validate_pdf_citations(pdf_bytes, ["Source 1"])
"""

from io import BytesIO

from docx import Document
from pypdf import PdfReader


def validate_docx_citations(docx_bytes: bytes, expected_citations: list[str]) -> bool:
    """Validate citations exist in DOCX content.

    Extracts all text from the DOCX document (paragraphs and tables)
    and checks if all expected citations are present.

    Args:
        docx_bytes: Raw bytes of the DOCX file
        expected_citations: List of citation strings to find

    Returns:
        bool: True if all expected citations are found
    """
    doc = Document(BytesIO(docx_bytes))

    # Collect all text from paragraphs
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    # Check all citations are present
    return all(citation in full_text for citation in expected_citations)


def validate_pdf_citations(pdf_bytes: bytes, expected_citations: list[str]) -> bool:
    """Validate citations exist in PDF content.

    Extracts text from all pages of the PDF and checks if all
    expected citations are present.

    Args:
        pdf_bytes: Raw bytes of the PDF file
        expected_citations: List of citation strings to find

    Returns:
        bool: True if all expected citations are found
    """
    pdf_reader = PdfReader(BytesIO(pdf_bytes))
    full_text = ""

    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            full_text += page_text + "\n"

    # Check all citations are present
    return all(citation in full_text for citation in expected_citations)


def validate_markdown_citations(
    markdown_content: str,
    expected_citations: list[str],
) -> bool:
    """Validate citations exist in Markdown content.

    Args:
        markdown_content: Markdown text content
        expected_citations: List of citation strings to find

    Returns:
        bool: True if all expected citations are found
    """
    return all(citation in markdown_content for citation in expected_citations)


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract all text content from a DOCX file.

    Args:
        docx_bytes: Raw bytes of the DOCX file

    Returns:
        str: Combined text from all paragraphs and tables
    """
    doc = Document(BytesIO(docx_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract all text content from a PDF file.

    Args:
        pdf_bytes: Raw bytes of the PDF file

    Returns:
        str: Combined text from all pages
    """
    pdf_reader = PdfReader(BytesIO(pdf_bytes))
    text_parts = []

    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def validate_export_format(
    content: bytes,
    expected_format: str,
) -> bool:
    """Validate exported content matches expected format.

    Checks file signatures/magic bytes for DOCX and PDF formats.

    Args:
        content: Raw bytes of the exported file
        expected_format: Expected format ("docx", "pdf", "markdown")

    Returns:
        bool: True if content matches expected format
    """
    if expected_format == "docx":
        # DOCX files are ZIP archives starting with PK
        return content[:4] == b"PK\x03\x04"
    elif expected_format == "pdf":
        # PDF files start with %PDF
        return content[:4] == b"%PDF"
    elif expected_format == "markdown":
        # Markdown is plain text - check it's valid UTF-8
        try:
            content.decode("utf-8")
            return True
        except UnicodeDecodeError:
            return False
    return False


def count_citations_in_content(
    content: str,
    citation_pattern: str | None = None,
) -> int:
    """Count citation markers in text content.

    Default pattern matches [1], [2], etc. citation markers.

    Args:
        content: Text content to search
        citation_pattern: Optional regex pattern for citations

    Returns:
        int: Number of citation markers found
    """
    import re

    pattern = citation_pattern or r"\[\d+\]"
    return len(re.findall(pattern, content))

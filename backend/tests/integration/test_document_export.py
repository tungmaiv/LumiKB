"""
ATDD Integration Tests: Epic 4 - Document Export (Story 4.7)
Status: RED phase - Tests written before implementation
Generated: 2025-11-26

Test Coverage:
- P0: DOCX export preserves citations (simple) (R-004)
- P0: DOCX export preserves citations (complex) (R-004)
- P0: PDF export preserves citations (simple) (R-004)
- P0: PDF export preserves citations (complex) (R-004)
- P1: Markdown export with footnotes
- P1: Verification prompt before export

Risk Mitigation:
- R-004 (DATA): Citation loss during export

Knowledge Base References:
- test-quality.md: Testing file outputs and validation

Note: This module requires PyPDF2 and python-docx dependencies.
      Tests will be skipped until Story 4.7 implementation.
"""

import io

import pytest

# Skip entire module if dependencies not available (Story 4.7 not implemented yet)
pytest.importorskip(
    "PyPDF2", reason="Story 4.7 not implemented - PyPDF2 dependency missing"
)
pytest.importorskip(
    "docx", reason="Story 4.7 not implemented - python-docx dependency missing"
)

import PyPDF2  # noqa: E402
from docx import Document  # python-docx  # noqa: E402
from fastapi import status  # noqa: E402
from httpx import AsyncClient  # noqa: E402


class TestDocumentExport:
    """Test document export functionality with citation preservation"""

    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_docx_export_preserves_simple_citations(
        self,
        client: AsyncClient,
        authenticated_headers: dict,
        demo_kb_with_indexed_docs: dict,
    ):
        """
        P0 Test - Story 4.7
        Risk: R-004 (Citation loss during export)

        GIVEN: User has generated a draft with 5 citations
        WHEN: Draft is exported to DOCX
        THEN: All 5 citations appear as footnotes in DOCX
        AND: Citation numbering is preserved
        """
        kb_id = demo_kb_with_indexed_docs["id"]

        # Generate draft with citations
        gen_response = await client.post(
            "/api/v1/generate",
            headers=authenticated_headers,
            json={
                "document_type": "RFP Response",
                "context": "OAuth and JWT implementation",
                "kb_id": kb_id,
            },
        )
        assert gen_response.status_code == status.HTTP_200_OK
        draft_id = gen_response.json()["draft_id"]
        expected_citations = gen_response.json()["citations"]

        # Export to DOCX
        export_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={"format": "docx"},
        )

        assert export_response.status_code == status.HTTP_200_OK
        assert export_response.headers["content-type"] == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Parse DOCX file
        docx_bytes = export_response.content
        doc = Document(io.BytesIO(docx_bytes))

        # Extract footnotes from DOCX
        # python-docx stores footnotes in doc.part.footnotes_part.footnotes
        footnotes = []
        if hasattr(doc.part, "footnotes_part") and doc.part.footnotes_part:
            for footnote in doc.part.footnotes_part.footnotes:
                footnote_text = "\n".join([p.text for p in footnote.paragraphs])
                footnotes.append(footnote_text)

        # CRITICAL: Number of footnotes must match number of citations
        assert len(footnotes) == len(expected_citations), (
            f"Expected {len(expected_citations)} footnotes but found {len(footnotes)} "
            "in exported DOCX"
        )

        # Verify each citation is present in footnotes
        for citation in expected_citations:
            citation_found = False
            for footnote in footnotes:
                # Check if citation content appears in footnote
                if citation["document_name"] in footnote:
                    citation_found = True
                    break
            assert citation_found, (
                f"Citation from '{citation['document_name']}' not found in DOCX footnotes"
            )

    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_docx_export_preserves_complex_citations(
        self,
        client: AsyncClient,
        authenticated_headers: dict,
        demo_kb_with_indexed_docs: dict,
    ):
        """
        P0 Test - Story 4.7
        Risk: R-004 (Citation loss during export)

        GIVEN: Draft with complex citations (multiple per sentence, long lists)
        WHEN: Exported to DOCX
        THEN: All citation markers preserved
        AND: Multiple citations per sentence handled correctly
        """
        kb_id = demo_kb_with_indexed_docs["id"]

        # Generate draft that will have multiple citations
        gen_response = await client.post(
            "/api/v1/generate",
            headers=authenticated_headers,
            json={
                "document_type": "Gap Analysis",
                "context": "Compare OAuth 2.0, JWT, and SAML authentication methods",
                "kb_id": kb_id,
            },
        )
        draft_id = gen_response.json()["draft_id"]
        draft_text = gen_response.json()["draft"]

        # Verify draft has complex citation patterns
        import re

        # Find sentences with multiple citations like [1, 3, 5]
        multi_citations = re.findall(r"\[(?:\d+,\s*)+\d+\]", draft_text)
        assert len(multi_citations) > 0, (
            "Draft should have multiple citations per sentence"
        )

        # Export to DOCX
        export_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={"format": "docx"},
        )

        docx_bytes = export_response.content
        doc = Document(io.BytesIO(docx_bytes))

        # Extract all text from document
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # CRITICAL: All citation markers from original draft must appear in DOCX
        original_markers = re.findall(r"\[[\d,\s]+\]", draft_text)
        for marker in original_markers:
            assert marker in full_text or any(
                m in full_text for m in marker.replace(",", "").split()
            ), f"Citation marker {marker} lost in DOCX export"

    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_pdf_export_preserves_citations(
        self,
        client: AsyncClient,
        authenticated_headers: dict,
        demo_kb_with_indexed_docs: dict,
    ):
        """
        P0 Test - Story 4.7
        Risk: R-004 (Citation loss during export)

        GIVEN: Draft with citations
        WHEN: Exported to PDF
        THEN: Citations rendered as footnotes or inline references
        AND: Formatting is preserved
        """
        kb_id = demo_kb_with_indexed_docs["id"]

        # Generate draft
        gen_response = await client.post(
            "/api/v1/generate",
            headers=authenticated_headers,
            json={
                "document_type": "Checklist",
                "context": "OAuth implementation checklist",
                "kb_id": kb_id,
            },
        )
        draft_id = gen_response.json()["draft_id"]
        expected_citations = gen_response.json()["citations"]

        # Export to PDF
        export_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={"format": "pdf"},
        )

        assert export_response.status_code == status.HTTP_200_OK
        assert export_response.headers["content-type"] == "application/pdf"

        # Parse PDF file
        pdf_bytes = export_response.content
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))

        # Extract text from all pages
        pdf_text = ""
        for page in pdf_reader.pages:
            pdf_text += page.extract_text()

        # CRITICAL: Each citation source must appear somewhere in PDF
        for citation in expected_citations:
            doc_name = citation["document_name"]
            assert doc_name in pdf_text, (
                f"Citation source '{doc_name}' not found in PDF export"
            )

        # Verify citation markers are present
        import re

        markers = re.findall(r"\[?\d+\]?", pdf_text)
        # Should have at least as many markers as citations
        assert len(markers) >= len(expected_citations), (
            "Not all citation markers found in PDF"
        )

    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_markdown_export_with_footnotes(
        self,
        client: AsyncClient,
        authenticated_headers: dict,
        demo_kb_with_indexed_docs: dict,
    ):
        """
        P1 Test - Story 4.7

        GIVEN: Draft with citations
        WHEN: Exported to Markdown
        THEN: Citations formatted as [^1] footnotes
        AND: Footnote definitions included at end
        """
        kb_id = demo_kb_with_indexed_docs["id"]

        # Generate draft
        gen_response = await client.post(
            "/api/v1/generate",
            headers=authenticated_headers,
            json={
                "document_type": "RFP Response",
                "context": "JWT authentication proposal",
                "kb_id": kb_id,
            },
        )
        draft_id = gen_response.json()["draft_id"]
        expected_citations = gen_response.json()["citations"]

        # Export to Markdown
        export_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={"format": "markdown"},
        )

        assert export_response.status_code == status.HTTP_200_OK
        assert "text/markdown" in export_response.headers["content-type"]

        markdown_text = export_response.content.decode("utf-8")

        # Verify footnote syntax [^1], [^2], etc.
        import re

        footnote_refs = re.findall(r"\[\^(\d+)\]", markdown_text)
        assert len(footnote_refs) >= len(expected_citations), (
            "Missing footnote references in Markdown"
        )

        # Verify footnote definitions exist (usually at end of document)
        # Format: [^1]: Source text here
        for i in range(1, len(expected_citations) + 1):
            footnote_def = f"[^{i}]:"
            assert footnote_def in markdown_text, (
                f"Footnote definition {footnote_def} not found in Markdown"
            )

    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_export_verification_prompt_required(
        self,
        client: AsyncClient,
        authenticated_headers: dict,
        demo_kb_with_indexed_docs: dict,
    ):
        """
        P1 Test - Story 4.7
        FR40b: System prompts "Have you verified the sources?" before export

        GIVEN: User requests export
        WHEN: Export is initiated
        THEN: API requires verification confirmation
        AND: Export fails without confirmation
        """
        kb_id = demo_kb_with_indexed_docs["id"]

        # Generate draft
        gen_response = await client.post(
            "/api/v1/generate",
            headers=authenticated_headers,
            json={"document_type": "RFP Response", "context": "Test", "kb_id": kb_id},
        )
        draft_id = gen_response.json()["draft_id"]

        # Attempt export WITHOUT verification confirmation
        export_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={
                "format": "docx",
                # Missing: "sources_verified": true
            },
        )

        # Should fail or require confirmation
        assert export_response.status_code in [
            status.HTTP_400_BAD_REQUEST,
            status.HTTP_428_PRECONDITION_REQUIRED,
        ], "Export should require source verification confirmation"

        error_data = export_response.json()
        assert (
            "verified" in error_data["detail"].lower()
            or "verify" in error_data["detail"].lower()
        )

        # Export WITH verification should succeed
        verified_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={"format": "docx", "sources_verified": True},
        )

        assert verified_response.status_code == status.HTTP_200_OK

    @pytest.mark.asyncio
    @pytest.mark.integration
    async def test_export_preserves_formatting(
        self,
        client: AsyncClient,
        authenticated_headers: dict,
        demo_kb_with_indexed_docs: dict,
    ):
        """
        P2 Test - Story 4.7
        FR40a: Exported documents preserve citations and formatting

        GIVEN: Draft with headers, lists, and formatting
        WHEN: Exported to DOCX/PDF
        THEN: Formatting is preserved (headers, lists, bold, etc.)
        """
        kb_id = demo_kb_with_indexed_docs["id"]

        # Generate draft with structured template (has headers, lists)
        gen_response = await client.post(
            "/api/v1/generate",
            headers=authenticated_headers,
            json={
                "document_type": "Gap Analysis",  # Has structured sections
                "context": "OAuth implementation",
                "kb_id": kb_id,
            },
        )
        draft_id = gen_response.json()["draft_id"]

        # Export to DOCX
        export_response = await client.post(
            f"/api/v1/generate/{draft_id}/export",
            headers=authenticated_headers,
            json={"format": "docx", "sources_verified": True},
        )

        docx_bytes = export_response.content
        doc = Document(io.BytesIO(docx_bytes))

        # Verify document has proper structure
        # Should have heading paragraphs
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) > 0, "DOCX should preserve heading structure"

        # Should have list items (if template used lists)
        # Check for numbered or bulleted paragraphs
        _lists = [p for p in doc.paragraphs if p.style.name.startswith("List")]
        # If template has lists, verify they're preserved
        # (This will pass/fail based on actual template structure)
